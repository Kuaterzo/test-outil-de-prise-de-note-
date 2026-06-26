"""Export des synthèses : Markdown (toujours), Word (.docx) et PDF (optionnels).

La synthèse produite est du Markdown simple (titres `##`/`###`, puces `-`/`*`,
gras `**…**`). Un petit analyseur dédié — volontairement minimal et testable —
convertit cette structure connue vers Word (`python-docx`) et PDF (`reportlab`),
deux bibliothèques pures Python qui ne nécessitent aucun outil système.

`python-docx` et `reportlab` sont importés paresseusement : le module reste
utilisable (et testable) même si elles ne sont pas installées.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

_SLUG_INVALID = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')
_SLUG_SPACES = re.compile(r"\s+")


class ExportError(RuntimeError):
    """Erreur lors de la génération d'un document (.docx / .pdf)."""


# ------------------------------------------------------------------ noms de fichiers
def slugify(text: str, max_length: int = 60) -> str:
    """Transforme un titre en composant de nom de fichier sûr (multi-OS)."""
    text = text.strip()
    text = _SLUG_INVALID.sub(" ", text)          # retire les caractères interdits
    text = _SLUG_SPACES.sub("_", text)           # espaces -> underscores
    text = text.strip("._")
    if len(text) > max_length:
        text = text[:max_length].rstrip("._")
    return text or "reunion"


def build_basename(title: str, when: datetime | None = None) -> str:
    """Construit un préfixe de fichier « AAAA-MM-JJ_HHhMM_Titre »."""
    when = when or datetime.now()
    return f"{when:%Y-%m-%d_%Hh%M}_{slugify(title)}"


# ------------------------------------------------------------------ analyse Markdown
def parse_inline(text: str) -> list[tuple[str, bool]]:
    """Découpe une ligne en fragments (texte, gras) selon les marqueurs `**`."""
    parts: list[tuple[str, bool]] = []
    for i, segment in enumerate(text.split("**")):
        if segment:
            parts.append((segment, i % 2 == 1))  # indices impairs = entre `**…**`
    return parts


def parse_markdown(markdown: str) -> list[tuple[str, str]]:
    """Convertit le Markdown de synthèse en blocs ``(type, texte)``.

    Types renvoyés : ``h1``, ``h2``, ``h3``, ``bullet``, ``para``.
    """
    blocks: list[tuple[str, str]] = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.lstrip()
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append(("bullet", stripped[2:].strip()))
        else:
            blocks.append(("para", line.strip()))
    return blocks


# ------------------------------------------------------------------ écriture
def _ensure_trailing_newline(text: str) -> str:
    return text if text.endswith("\n") else text + "\n"


def save_synthesis(
    synthesis: str,
    output_dir: Path,
    title: str,
    *,
    transcript: str | None = None,
    when: datetime | None = None,
) -> dict[str, Path]:
    """Écrit la synthèse Markdown (et éventuellement la transcription).

    Renvoie un dictionnaire des chemins créés (clés ``synthesis`` et,
    le cas échéant, ``transcript``). Les formats Word/PDF sont produits
    séparément via :func:`render_docx` / :func:`render_pdf`.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    basename = build_basename(title, when)
    paths: dict[str, Path] = {}

    synthesis_path = output_dir / f"{basename}.md"
    synthesis_path.write_text(_ensure_trailing_newline(synthesis), encoding="utf-8")
    paths["synthesis"] = synthesis_path

    if transcript is not None:
        transcript_path = output_dir / f"{basename}_transcription.txt"
        transcript_path.write_text(_ensure_trailing_newline(transcript), encoding="utf-8")
        paths["transcript"] = transcript_path

    return paths


def render_docx(synthesis: str, path: Path, title: str) -> Path:
    """Génère un document Word (.docx) à partir de la synthèse Markdown."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dépendance optionnelle
        raise ExportError(
            "Le paquet « python-docx » est requis pour l'export Word "
            "(pip install python-docx)."
        ) from exc

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(title, level=0)

    for kind, text in parse_markdown(synthesis):
        if kind in ("h1", "h2"):
            document.add_heading(text, level=1)
        elif kind == "h3":
            document.add_heading(text, level=2)
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            _add_docx_runs(paragraph, text)
        else:
            paragraph = document.add_paragraph()
            _add_docx_runs(paragraph, text)

    document.save(str(path))
    return path


def _add_docx_runs(paragraph, text: str) -> None:
    """Ajoute les fragments d'une ligne à un paragraphe Word, gras compris."""
    for segment, bold in parse_inline(text):
        run = paragraph.add_run(segment)
        run.bold = bold


def render_pdf(synthesis: str, path: Path, title: str) -> Path:
    """Génère un document PDF à partir de la synthèse Markdown (via reportlab)."""
    try:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:  # pragma: no cover - dépendance optionnelle
        raise ExportError(
            "Le paquet « reportlab » est requis pour l'export PDF "
            "(pip install reportlab)."
        ) from exc

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    bullet_style = ParagraphStyle(
        "Bullet", parent=styles["BodyText"], leftIndent=18, alignment=TA_LEFT, spaceAfter=4
    )
    story = [Paragraph(_inline_to_html(title), styles["Title"]), Spacer(1, 8)]

    for kind, text in parse_markdown(synthesis):
        html = _inline_to_html(text)
        if kind in ("h1", "h2"):
            story.append(Spacer(1, 6))
            story.append(Paragraph(html, styles["Heading2"]))
        elif kind == "h3":
            story.append(Paragraph(html, styles["Heading3"]))
        elif kind == "bullet":
            story.append(Paragraph(f"•&nbsp;{html}", bullet_style))
        else:
            story.append(Paragraph(html, styles["BodyText"]))
            story.append(Spacer(1, 4))

    document = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
        title=title,
    )
    document.build(story)
    return path


def _inline_to_html(text: str) -> str:
    """Convertit le gras Markdown en balisage minimal pour reportlab (XML-safe)."""
    from xml.sax.saxutils import escape

    out = []
    for segment, bold in parse_inline(text):
        safe = escape(segment)
        out.append(f"<b>{safe}</b>" if bold else safe)
    return "".join(out)


__all__ = [
    "ExportError",
    "slugify",
    "build_basename",
    "parse_inline",
    "parse_markdown",
    "save_synthesis",
    "render_docx",
    "render_pdf",
]
